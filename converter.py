import csv
import unicodedata
import re
import os
import json
from docx import Document
from enum import Enum
import pandas as pd
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

# Globals
SETTINGS_PATH = "settings.json"

def get_docs_files(pathToDocsFiles):
    # List to store the filenames of Word documents
    fileNames = []

    # Iterate through all files in the directory
    for filename in os.listdir(pathToDocsFiles):
        if filename.endswith(".docx"):
            # Add the filename to the list
            fileNames.append(filename)
    
    return fileNames

def read_settings():
    with open(SETTINGS_PATH, "r") as json_file:
        data = json.load(json_file)
    
    return data

def to_excel_table(excelPath, df, index=False):
    writer = pd.ExcelWriter(excelPath, engine='openpyxl', mode='a', if_sheet_exists='overlay')
    sheetName = "Evaluation"
    tableName = "T_Evaluation"
    df.to_excel(writer, sheet_name=sheetName , startrow=0, header=True, index=index)
    
    # Get the xlsxwriter workbook and worksheet objects.
    worksheet = writer.sheets[sheetName]
    table = worksheet.tables[tableName]
    
    # Get the dimensions of the dataframe.
    (max_row, max_col) = df.shape
    
    if index:
        max_col += 1
    
    # When A-Z
    if max_col <= 24:
        maxColString = chr(max_col+64)
    else:
        # When over Z
        maxColString = 'A' + chr(64 + max_col%26)

    ref = 'A1:{0}{1}'.format(maxColString, str(max_row +1))

    del worksheet.tables[tableName]

    tab = Table(displayName=tableName, ref=ref)

    style = TableStyleInfo(name="TableStyleLight15", showFirstColumn=False,
                    showLastColumn=False, showRowStripes=True, showColumnStripes=True)
    tab.tableStyleInfo = style

    #table.ref = ref
    worksheet.add_table(tab)

    for column_cells in worksheet.columns:
        length = max(len(str(cell.value)) for cell in column_cells)
        worksheet.column_dimensions[column_cells[0].column_letter].width = length

    writer.close()
    
    

class Categories(Enum):
    CHILD = 1
    PARENT = 2
    UNDEFINED = 3

class ClassifiedParagraph():
    FILTER_EXEPTIONS = read_settings()["notIgnoredInSquareBrackets"]
    
    def __init__(self, paraText:str):
        self.category = self.classify(paraText)
        if self.category != Categories.UNDEFINED:
            self.text = paraText[2:].strip()
        else:
            self.text = paraText
        self.numberWords = self.count_words()
        self.numberLetters = self.count_chars()

    def classify(self, paraText):
        first_letters = paraText[0:3].lower()
        if 'c:' in first_letters:
            return Categories.CHILD
        
        if 'p:' in first_letters:
            return Categories.PARENT
        
        return Categories.UNDEFINED
    
    def get_clean_text(self):
        # Filter [*Emotion*] from string
        
        exceptions_pattern =  "?!" + '|'.join(f"{re.escape(exception)}" for exception in self.FILTER_EXEPTIONS)
        pattern = re.compile(r'\[((' + exceptions_pattern + r').*?)\]')
        clean_text =  re.sub(pattern, '', self.text.lower())
        badSpacingPattern = r'\s[.!?]'
        clean_text = re.sub(badSpacingPattern, '', clean_text)

        return clean_text
        
    def count_words(self):
        # Remove special characters from text
        return len(self.get_clean_text().split())
    
    def count_chars(self):
        # count letters from A to Z
        letters = re.findall(r'[a-zA-Z0-9]', self.get_clean_text())
        letters_string = ''.join(letters)
        return len(letters)

class AnalyseDocument():
    classifiedParagraphs = []

    def __init__(self, document:Document, id) -> None:

        self.id = id
        self.wordsParent = 0
        self.wordsChild = 0
        self.numberParagraphs = 0
        self.numberParentParagraphs = 0
        self.numberChildParagraphs = 0
        self.lettersParent = 0
        self.lettersChild = 0

        for paragraph in document.paragraphs:
            # categorise sentence
            text = paragraph.text.strip()
            if text != '':
                paragraph = ClassifiedParagraph(text)
                self.classifiedParagraphs.append(paragraph)

                if paragraph.category == Categories.PARENT:
                    self.numberParentParagraphs += 1
                    self.numberParagraphs += 1
                    self.wordsParent += paragraph.numberWords
                    self.lettersParent += paragraph.numberLetters
                
                elif paragraph.category == Categories.CHILD:
                    self.numberChildParagraphs += 1
                    self.numberParagraphs += 1
                    self.wordsChild += paragraph.numberWords
                    self.lettersChild += paragraph.numberLetters
    
    def to_csv(self, file):
        with open(file, mode='w', newline='\n') as f:
            header = ['Category', 'Text', 'Words' , 'Letters', 'Open Questions', 'Closed Questions']
            csvWriter = csv.writer(f, delimiter=';', quotechar='"', quoting=csv.QUOTE_NONNUMERIC)

            csvWriter.writerow(header)

            for cp in self.classifiedParagraphs:
                csvWriter.writerow([cp.category.name, cp.text, cp.numberWords, cp.numberLetters, None, None])
    
    def get_table_data(self):
        return {
            'Chiffre' : self.id,
            # 'Total parent questions': None,
            # 'Parent questions': None,
            # 'Open ended parent questions': None,
            # 'Number of parent utterances': None,
            # 'Mean length of parent utterances': None,
            'Words parent': self.wordsParent,
            # 'Total child questions': None,
            # 'Child questions': None,
            # 'Number of child utterances': None,
            # 'Open ended child questions': None,
            # 'Mean length of child utterances': None,
            'Words child': self.wordsChild
        } 


def run():
    print("Start run")
    SETTINGS = read_settings()
    docFileNames = get_docs_files(SETTINGS['docFilePath'])
    print(str(docFileNames))

    dfOrigin = pd.read_excel(SETTINGS['pathToExcelFile'], sheet_name='Evaluation')
    read_data = []

    for docFileName in docFileNames:
        print("Read " + docFileName)
        document = Document(SETTINGS['docFilePath'] + docFileName) 
        csvFileName = docFileName.replace(".docx", ".csv")
        analysed = AnalyseDocument(document, docFileName.replace(".docx", ""))
        print("Create csv: " + csvFileName)
        analysed.to_csv(SETTINGS['csvFiles'] + csvFileName)
        read_data.append(analysed.get_table_data())
    
    dfNew = pd.DataFrame(read_data)

    # Append new Inputs to origin. Delete duplicates and keep the origin
    df = pd.concat([dfOrigin, dfNew], axis=0).drop_duplicates(subset=['Chiffre'], keep='first')
    print(str(df))
    to_excel_table(SETTINGS['pathToExcelFile'], df, False)
    print("End run")

run()

